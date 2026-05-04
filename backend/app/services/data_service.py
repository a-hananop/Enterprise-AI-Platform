"""
Data Service - File parsing, metadata extraction, statistics
Handles: CSV, Excel, JSON, PDF, DOCX, TXT
"""
import os
import json
from typing import Dict, List, Optional, Any
import pandas as pd
import numpy as np


class DataService:

    def _get_ext(self, source_type: Any) -> str:
        if hasattr(source_type, 'value'):
            return str(source_type.value).lower()
        return str(source_type).split('.')[-1].lower()

    async def extract_metadata(self, file_path: str, source_type: str) -> Dict:
        """Extract metadata from uploaded file"""
        ext = self._get_ext(source_type)
        try:
            if ext in ["csv", "excel", "json"]:
                return await self._tabular_metadata(file_path, ext)
            elif ext in ["pdf"]:
                return await self._pdf_metadata(file_path)
            elif ext in ["docx"]:
                return await self._docx_metadata(file_path)
            elif ext in ["txt"]:
                return await self._txt_metadata(file_path)
        except Exception as e:
            print(f"Metadata extraction error: {e}")
        return {}

    async def _tabular_metadata(self, file_path: str, source_type: str) -> Dict:
        df = self._load_df(file_path, source_type)
        if df is None:
            return {}

        columns_info = []
        for col in df.columns:
            dtype = str(df[col].dtype)
            null_count = int(df[col].isnull().sum())
            unique_count = int(df[col].nunique())
            col_info = {
                "name": col,
                "dtype": dtype,
                "null_count": null_count,
                "null_percent": round(null_count / len(df) * 100, 1) if len(df) > 0 else 0,
                "unique_count": unique_count,
            }
            if df[col].dtype in [np.float64, np.int64, np.float32, np.int32]:
                col_info.update({
                    "min": self._safe_float(df[col].min()),
                    "max": self._safe_float(df[col].max()),
                    "mean": self._safe_float(df[col].mean()),
                    "std": self._safe_float(df[col].std()),
                })
            elif df[col].dtype == object:
                top_vals = df[col].value_counts().head(5).to_dict()
                col_info["top_values"] = {str(k): int(v) for k, v in top_vals.items()}
            columns_info.append(col_info)

        # Safe preview (handle NaN values)
        preview_df = df.head(5).copy()
        preview_df = preview_df.replace({np.nan: None, np.inf: None, -np.inf: None})
        
        def convert_value(v):
            if v is None:
                return None
            if isinstance(v, (np.integer,)):
                return int(v)
            if isinstance(v, (np.floating,)):
                return float(v) if not np.isnan(v) else None
            return v
        
        preview_data = []
        for record in preview_df.to_dict("records"):
            preview_data.append({k: convert_value(v) for k, v in record.items()})

        return {
            "row_count": len(df),
            "column_count": len(df.columns),
            "columns_info": columns_info,
            "preview_data": preview_data,
        }

    async def _pdf_metadata(self, file_path: str) -> Dict:
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                text_sample = ""
                for page in pdf.pages[:2]:
                    t = page.extract_text()
                    if t:
                        text_sample += t[:200]
            return {
                "row_count": page_count,
                "column_count": 1,
                "preview_data": [{"page": 1, "content_preview": text_sample[:500]}],
            }
        except Exception:
            return {"row_count": 1, "column_count": 1}

    async def _docx_metadata(self, file_path: str) -> Dict:
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return {
                "row_count": len(paragraphs),
                "column_count": 1,
                "preview_data": [{"paragraph": i+1, "text": p[:200]} for i, p in enumerate(paragraphs[:5])],
            }
        except Exception:
            return {"row_count": 1, "column_count": 1}

    async def _txt_metadata(self, file_path: str) -> Dict:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        lines = content.split("\n")
        return {
            "row_count": len(lines),
            "column_count": 1,
            "preview_data": [{"line": i+1, "content": l[:200]} for i, l in enumerate(lines[:5])],
        }

    async def preview(self, file_path: str, source_type: str, rows: int = 10) -> List[Dict]:
        """Return preview rows"""
        ext = self._get_ext(source_type)
        try:
            if ext in ["csv", "excel", "json"]:
                df = self._load_df(file_path, ext)
                if df is not None:
                    preview_df = df.head(rows).replace({np.nan: None, np.inf: None, -np.inf: None})
                    
                    def convert_value(v):
                        if v is None:
                            return None
                        if isinstance(v, (np.integer,)):
                            return int(v)
                        if isinstance(v, (np.floating,)):
                            return float(v) if not np.isnan(v) else None
                        return v
                    
                    result = []
                    for record in preview_df.to_dict("records"):
                        result.append({k: convert_value(v) for k, v in record.items()})
                    return result
            elif ext == "pdf":
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    pages = []
                    for i, page in enumerate(pdf.pages[:rows]):
                        t = page.extract_text()
                        pages.append({"page": i+1, "content": t[:500] if t else ""})
                return pages
            elif ext == "docx":
                from docx import Document
                doc = Document(file_path)
                paras = [{"paragraph": i+1, "text": p.text} for i, p in enumerate(doc.paragraphs[:rows]) if p.text.strip()]
                return paras
            elif ext == "txt":
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    lines = f.readlines()[:rows]
                return [{"line": i+1, "content": l.strip()} for i, l in enumerate(lines)]
        except Exception as e:
            print(f"Preview error: {e}")
        return []

    async def compute_stats(self, file_path: str, source_type: str) -> Dict:
        """Compute statistical summary of tabular data"""
        df = self._load_df(file_path, self._get_ext(source_type))
        if df is None:
            return {}

        numeric_cols = df.select_dtypes(include=[np.number]).columns.tolist()
        categorical_cols = df.select_dtypes(include=["object"]).columns.tolist()

        stats = {
            "shape": {"rows": len(df), "columns": len(df.columns)},
            "missing_values": {col: int(df[col].isnull().sum()) for col in df.columns},
            "numeric_summary": {},
            "categorical_summary": {},
            "correlations": {},
        }

        if numeric_cols:
            desc = df[numeric_cols].describe().to_dict()
            stats["numeric_summary"] = {
                col: {k: round(float(v), 4) if not np.isnan(v) else None for k, v in vals.items()}
                for col, vals in desc.items()
            }
            if len(numeric_cols) > 1:
                corr = df[numeric_cols].corr()
                stats["correlations"] = {
                    col: {c: round(float(v), 4) for c, v in row.items() if not np.isnan(v)}
                    for col, row in corr.to_dict().items()
                }

        for col in categorical_cols[:5]:
            vc = df[col].value_counts().head(10)
            stats["categorical_summary"][col] = {str(k): int(v) for k, v in vc.items()}

        return stats

    async def clean_data(self, file_path: str, source_type: str, operations: Dict) -> Dict:
        """Apply cleaning operations"""
        df = self._load_df(file_path, self._get_ext(source_type))
        if df is None:
            return {"error": "Cannot load file"}

        changes = []

        if operations.get("drop_duplicates"):
            before = len(df)
            df = df.drop_duplicates()
            changes.append(f"Removed {before - len(df)} duplicate rows")

        if operations.get("fill_nulls"):
            strategy = operations["fill_nulls"]
            for col in df.select_dtypes(include=[np.number]).columns:
                if strategy == "mean":
                    df[col] = df[col].fillna(df[col].mean())
                elif strategy == "median":
                    df[col] = df[col].fillna(df[col].median())
                elif strategy == "zero":
                    df[col] = df[col].fillna(0)
            changes.append(f"Filled null values using {strategy}")

        if operations.get("drop_columns"):
            cols = operations["drop_columns"]
            df = df.drop(columns=[c for c in cols if c in df.columns], errors="ignore")
            changes.append(f"Dropped columns: {cols}")

        # Save cleaned file
        clean_path = file_path.replace(".", "_cleaned.")
        if self._get_ext(source_type) == "csv":
            df.to_csv(clean_path, index=False)
        elif self._get_ext(source_type) == "excel":
            df.to_excel(clean_path, index=False)

        return {
            "changes": changes,
            "rows": len(df),
            "columns": len(df.columns),
            "cleaned_file": os.path.basename(clean_path),
        }

    def _load_df(self, file_path: str, source_type: str) -> Optional[pd.DataFrame]:
        try:
            if source_type in ["csv"]:
                return pd.read_csv(file_path)
            elif source_type in ["excel"]:
                return pd.read_excel(file_path)
            elif source_type in ["json"]:
                return pd.read_json(file_path)
        except Exception as e:
            print(f"DataFrame load error: {e}")
        return None

    def _safe_float(self, val) -> Optional[float]:
        try:
            v = float(val)
            return round(v, 4) if not np.isnan(v) and not np.isinf(v) else None
        except Exception:
            return None
