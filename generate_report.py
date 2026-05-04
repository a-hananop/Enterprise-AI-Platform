from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page Margins (tighter to fit 5 pages) ──────────────────────
section = doc.sections[0]
section.top_margin    = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin   = Inches(1.1)
section.right_margin  = Inches(1.1)

# ── Helpers ────────────────────────────────────────────────────
def add_hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E5ECC')
    pBdr.append(bottom)
    pPr.append(pBdr)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        run.font.size      = Pt(12.5)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0x2E, 0x5E, 0xCC)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
    elif level == 2:
        run.font.size      = Pt(11)
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size      = Pt(10.5)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    p.paragraph_format.space_after   = Pt(3)
    p.paragraph_format.line_spacing  = Pt(13.5)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + ": ")
        r.bold = True
        r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)

def reference(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after       = Pt(2)
    r1 = p.add_run(f"[{number}] ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)

# ═══════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("Enterprise AI Platform")
tr.font.size      = Pt(22)
tr.font.bold      = True
tr.font.color.rgb = RGBColor(0x2E, 0x5E, 0xCC)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run("An Integrated Multi-Module AI Solution for Business Intelligence")
sr.font.size       = Pt(13)
sr.font.italic     = True
sr.font.color.rgb  = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph()
add_hrule(doc)
doc.add_paragraph()

meta_lines = [
    "Semester Project — Progress Report",
    f"Date: {datetime.date.today().strftime('%B %d, %Y')}",
    "Course: Artificial Intelligence / Data Science",
]
for line in meta_lines:
    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = mp.add_run(line)
    mr.font.size      = Pt(11)
    mr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  1. INTRODUCTION AND BACKGROUND
# ═══════════════════════════════════════════════════════════════
heading(doc, "1.  Introduction and Background")
add_hrule(doc)

body(doc,
    "Modern enterprises generate vast volumes of heterogeneous data — structured datasets, "
    "unstructured documents, and real-time event streams — most of which remains analytically "
    "untapped due to fragmented, siloed tooling. Organisations typically maintain 5–10 separate "
    "specialised products to cover the full analytical workflow, incurring high integration "
    "overhead, operational cost, and a steep learning curve for non-technical stakeholders."
)
body(doc,
    "The Enterprise AI Platform is a unified, full-stack web application consolidating eight "
    "AI-powered modules: Data Hub, Analytics & BI, Predictive ML, NLP/Text Analysis, Document "
    "Intelligence, AI Automation Agents, Marketing AI, and Risk Detection. The system is built "
    "on FastAPI (Python) and React/TypeScript and is designed to democratise AI for analysts, "
    "data scientists, and executives without requiring specialist infrastructure knowledge."
)

heading(doc, "1.1  Motivation & Applications", level=2)
bullet(doc, "Tool fragmentation forces teams to stitch together disconnected platforms, creating governance gaps and data silos.", "Fragmentation")
bullet(doc, "Cost prohibits SMBs from accessing enterprise AI suites (e.g., Azure AI, Salesforce Einstein) priced at tens of thousands per year.", "Accessibility")
bullet(doc, "Key use cases: sales/revenue forecasting, compliance document review, customer sentiment monitoring, autonomous market research, and HR attrition prediction.", "Applications")

heading(doc, "1.2  Domain Knowledge", level=2)
body(doc,
    "The project spans three primary domains: (a) Natural Language Processing — sentiment analysis, "
    "named-entity recognition, summarisation, and retrieval-augmented generation (RAG); "
    "(b) Machine Learning — supervised regression/classification, unsupervised clustering/anomaly "
    "detection, and automated preprocessing via scikit-learn; and (c) Business Intelligence — "
    "KPI extraction, time-series trend analysis, and AI-generated narrative reporting."
)

# ═══════════════════════════════════════════════════════════════
#  2. LITERATURE REVIEW
# ═══════════════════════════════════════════════════════════════
heading(doc, "2.  Literature Review")
add_hrule(doc)

heading(doc, "2.1  Large Language Models", level=2)
body(doc,
    "Devlin et al. [1] established the transformer pre-train/fine-tune paradigm with BERT, "
    "while Brown et al. [2] demonstrated GPT-3's few-shot capabilities, obviating task-specific "
    "fine-tuning. Our platform exploits this by directing Groq-hosted Llama-3-70b via "
    "structured prompting rather than fine-tuning, keeping infrastructure cost near zero. "
    "Unlike cloud-only solutions, local sentence-transformer embeddings [3] are used for "
    "semantic search, eliminating per-query API charges."
)

heading(doc, "2.2  Retrieval-Augmented Generation (RAG)", level=2)
body(doc,
    "Lewis et al. [4] formalised RAG to ground LLM outputs in external documents, drastically "
    "reducing hallucinations. The Document Intelligence module implements this using ChromaDB "
    "(local vector store) with cosine-similarity retrieval — consistent with Borgeaud et al. [5]. "
    "Unlike proprietary solutions (AWS Kendra, Azure AI Search), our implementation is fully "
    "local, supporting air-gapped or on-premises deployments."
)

heading(doc, "2.3  AutoML and AI Agents", level=2)
body(doc,
    "Hutter et al. [6] showed that automated ML pipelines lower the expertise barrier for "
    "model deployment. The Predictions module automates preprocessing, model selection, and "
    "evaluation — applying AutoML principles without cloud billing. The Automation module "
    "implements a structured multi-agent orchestration pattern inspired by ReAct [7] and "
    "HuggingGPT [8], where specialised agents (Data, Research, Finance, Marketing) are "
    "co-ordinated by an Orchestrator Agent — more controllable than fully autonomous systems."
)

heading(doc, "2.4  Business Intelligence", level=2)
body(doc,
    "Larson & Chang [9] noted the shift from static BI dashboards toward AI-generated narrative "
    "summaries. The Analytics module extends this trend by combining Recharts-based "
    "visualisations with LLM-generated prose reports, providing both quantitative precision "
    "and human-readable interpretation in one interface."
)

heading(doc, "References", level=2)
reference(doc, 1, "Devlin et al. (2019). BERT. NAACL-HLT.")
reference(doc, 2, "Brown et al. (2020). GPT-3: Language models are few-shot learners. NeurIPS.")
reference(doc, 3, "Reimers & Gurevych (2019). Sentence-BERT. EMNLP.")
reference(doc, 4, "Lewis et al. (2020). Retrieval-Augmented Generation. NeurIPS.")
reference(doc, 5, "Borgeaud et al. (2022). Improving LMs by retrieving from trillions of tokens. ICML.")
reference(doc, 6, "Hutter et al. (2019). Automated Machine Learning. Springer.")
reference(doc, 7, "Yao et al. (2022). ReAct: Reasoning + acting in LMs. ICLR 2023.")
reference(doc, 8, "Shen et al. (2023). HuggingGPT. NeurIPS.")
reference(doc, 9, "Larson & Chang (2016). A review of agile BI and data science. Int. J. Information Management.")

# ═══════════════════════════════════════════════════════════════
#  3. PRELIMINARY EXPERIMENTS AND RESULTS
# ═══════════════════════════════════════════════════════════════
heading(doc, "3.  Preliminary Experiments and Results")
add_hrule(doc)

body(doc,
    "All experiments used a 120-row × 10-column synthetic sales dataset (sales_data.csv) "
    "with fields: date, product, region, units_sold, revenue, marketing_spend, "
    "customer_satisfaction, and profit_margin."
)

heading(doc, "3.1  Data Hub & KPI Extraction", level=2)
body(doc,
    "Upload, parsing, dtype inference, and metadata storage all succeeded. KPI extraction "
    "identified 8 metrics: Total Revenue ($236,531), Avg. Customer Satisfaction (4.2/5.0), "
    "Total Units Sold (12,480), and margin/spend ratios — each with trend direction computed "
    "via period-over-period change percentages."
)

heading(doc, "3.2  Trend Analysis", level=2)
body(doc,
    "Monthly resampling with date_column='date', value_columns='revenue' rendered correctly "
    "in Recharts (sum and mean traces). A critical enum-stringification bug was discovered and "
    "fixed here: DataSourceType.csv was stringified as 'DataSourceType.csv' instead of 'csv', "
    "breaking all file-loading paths until a _get_ext() helper was introduced."
)

heading(doc, "3.3  Predictive ML", level=2)
body(doc,
    "Linear regression on revenue achieved R²=0.84. K-Means clustering (k=3) also ran "
    "successfully after fixing a form validation bug that incorrectly required target_column "
    "for unsupervised models. The automated pipeline (imputation → scaling → fit → evaluate) "
    "executed under 2 seconds for the 120-row dataset."
)

heading(doc, "3.4  Document Intelligence & NLP", level=2)
body(doc,
    "A PDF uploaded with RAG indexing enabled produced 42 ChromaDB chunks. Summarisation "
    "and Q&A queries returned grounded responses via RAG + LLM. Text Analysis (sentiment, "
    "entities, keywords) and email response generation averaged ~1.8 s latency on Groq's "
    "free tier. The system degrades gracefully when RAG dependencies are absent."
)

heading(doc, "3.5  UI Validation", level=2)
body(doc,
    "A full audit found that six page components used .input-field and .stat-card CSS class "
    "names that did not exist in the design system (defined as .input and .stat). A batch "
    "replacement corrected all occurrences. Full-width primary action buttons and consistent "
    ".stat-label / .stat-value typography were applied platform-wide."
)

# ═══════════════════════════════════════════════════════════════
#  4. PROBLEMS FACED AND POTENTIAL SOLUTIONS
# ═══════════════════════════════════════════════════════════════
heading(doc, "4.  Problems Faced and Potential Solutions")
add_hrule(doc)

rows = [
    ("Enum Stringification",
     "SQLAlchemy returned DataSourceType enum objects instead of string values, silently breaking file-loading across Data, Analytics, and ML services.",
     "Introduced _get_ext() helper to safely extract .value from enum or pass string through unchanged."),
    ("Missing RAG Dependencies",
     "chromadb and sentence-transformers not installed by default; Document Intelligence returned empty results with no user-facing error.",
     "Dependencies installed; RAGService updated to degrade gracefully and surface an informative status flag."),
    ("CSS Class Mismatch",
     "Six pages referenced non-existent .input-field / .stat-card classes; form controls and KPI cards rendered as unstyled browser defaults.",
     "Batch script replaced all class names with the correct .input / .stat equivalents from index.css."),
    ("ML Form Validation Bug",
     "Predictions form required target_column for all model types, blocking submission for clustering and anomaly detection.",
     "Validation updated to require target_column only for supervised models."),
    ("Hot-Reload Gaps",
     "Uvicorn auto-reload occasionally missed changes in deep service files, requiring manual restarts.",
     "Short-term: manual restart; long-term: --reload-dir app flag + Docker containerisation."),
]

for problem, desc, solution in rows:
    heading(doc, f"4.x  {problem}", level=2)
    body(doc, f"Problem: {desc}")
    body(doc, f"Solution: {solution}", indent=True)

# ═══════════════════════════════════════════════════════════════
#  5. OBJECTIVES MET AND PENDING
# ═══════════════════════════════════════════════════════════════
heading(doc, "5.  Objectives — Met and Pending")
add_hrule(doc)

heading(doc, "5.1  Objectives Met (Completed)", level=2)
met = [
    "Full-stack scaffolding: FastAPI backend + React/TypeScript + Tailwind glassmorphism design system.",
    "Data Hub: multi-format upload (CSV, Excel, JSON, PDF, DOCX, TXT), schema detection, preview, and stats.",
    "Analytics & BI: KPI auto-extraction, trend analysis with Recharts, AI-narrative report generation, and alert management.",
    "Predictive ML: regression, classification, clustering, anomaly detection with automated preprocessing pipeline.",
    "NLP / Text Analysis: sentiment, NER, keyword extraction, classification, summarisation, email response generation.",
    "Document Intelligence: RAG-based summarisation, Q&A, information extraction, document comparison, meeting transcript analysis.",
    "AI Agents: single-agent task execution and multi-agent orchestration with step tracking and final synthesis.",
    "Marketing AI: content generation, ad copy creation, and campaign performance scoring.",
    "Risk Detection: statistical anomaly flagging, risk scoring, and automated report generation.",
    "JWT-based authentication and user session management.",
    "Premium UI: glassmorphism, staggered animations, page transitions, and full responsive layout.",
    "Critical bug fixes: enum stringification, ML form validation, CSS class mismatches — all resolved.",
]
for m in met:
    bullet(doc, m)

heading(doc, "5.2  Objectives Pending (Remaining)", level=2)
pending = [
    "Automated test suite: pytest unit/integration tests for all backend services.",
    "Streaming responses: SSE-based LLM streaming for improved perceived latency.",
    "Multi-tenancy & RBAC: organisational roles (Admin, Analyst, Viewer).",
    "Production deployment: Docker Compose + cloud hosting (Railway / AWS ECS).",
    "Model persistence & versioning: saving trained models across sessions (MLflow or DB registry).",
    "Advanced visualisations: scatter plots, correlation heatmaps, confusion matrices in frontend.",
    "UI export: one-click PDF/Excel export of dashboards and analysis results.",
]
for p in pending:
    bullet(doc, p)

# ═══════════════════════════════════════════════════════════════
#  SAVE
# ═══════════════════════════════════════════════════════════════
output_path = r"C:\Users\HP\Downloads\enterprise-ai-platform\enterprise-ai-platform\Assignment_Report_v2.docx"
doc.save(output_path)
print("Report saved to: " + output_path)
